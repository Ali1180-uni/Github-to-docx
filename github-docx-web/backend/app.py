"""
GitHub to DOCX Web API - Backend Server
Flask-based REST API for converting GitHub repositories to Word documents
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import requests
import re
import datetime
import uuid
import shutil
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from urllib.parse import urlparse
import threading
import time

app = Flask(__name__)
CORS(app)

# Store job status
jobs = {}

class GitHubFolderAgent:
    def __init__(self, output_directory="Output_Reports", github_token=None, job_id=None):
        self.output_directory = output_directory
        self.github_token = github_token
        self.folder_documents = {}
        self.source_url = ""
        self.job_id = job_id
        self.processed_files = 0
        self.total_files = 0
        self.current_file = ""
        self.status = "initializing"
        self.error = None
        
        if not os.path.exists(self.output_directory):
            os.makedirs(self.output_directory)

    def _setup_document_style(self, doc):
        """Sets up a professional visual style for the Word document."""
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Segoe UI'
        font.size = Pt(10.5)
        
        h1 = doc.styles['Heading 1']
        h1.font.name = 'Segoe UI'
        h1.font.size = Pt(14)
        h1.font.color.rgb = RGBColor(31, 78, 121)
        h1.font.bold = True
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(3)

    def _set_paragraph_shading(self, paragraph, color_hex):
        """Adds background color/shading to a paragraph."""
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        paragraph.paragraph_format.element.get_or_add_pPr().append(shd)

    def get_headers(self):
        """Returns headers for GitHub API requests."""
        headers = {'Accept': 'application/vnd.github.v3+json'}
        if self.github_token:
            headers['Authorization'] = f'token {self.github_token}'
        return headers

    def parse_github_url(self, url):
        """Converts a browser URL to a GitHub API URL."""
        parsed = urlparse(url)
        path_parts = parsed.path.strip('/').split('/')
        
        if len(path_parts) < 2:
            raise ValueError("Invalid GitHub URL")

        owner = path_parts[0]
        repo = path_parts[1]
        
        api_url = f"https://api.github.com/repos/{owner}/{repo}/contents"
        params = {}

        if len(path_parts) > 3 and path_parts[2] == 'tree':
            branch = path_parts[3]
            subpath = "/".join(path_parts[4:])
            api_url = f"{api_url}/{subpath}"
            params['ref'] = branch
            
        return api_url, params

    def natural_key(self, string_):
        """Returns a key for natural sorting."""
        return [int(s) if s.isdigit() else s.lower() for s in re.split(r'(\d+)', string_)]

    def sanitize_filename(self, name):
        """Sanitizes folder name to be used as a valid filename."""
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            name = name.replace(char, '_')
        return name.strip()

    def get_or_create_document(self, folder_name):
        """Gets existing document for folder or creates a new one."""
        if folder_name not in self.folder_documents:
            doc = Document()
            self._setup_document_style(doc)
            
            title = doc.add_heading(f'{folder_name}', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title.runs[0].font.name = 'Segoe UI'
            title.runs[0].font.color.rgb = RGBColor(16, 54, 92)

            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.add_run(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n").bold = True
            p.add_run(f"Source: {self.source_url}")
            doc.add_page_break()
            
            self.folder_documents[folder_name] = doc
            
        return self.folder_documents[folder_name]

    def get_folder_name_from_path(self, file_path):
        """Extracts the immediate parent folder name from a file path."""
        path_parts = file_path.split('/')
        if len(path_parts) > 1:
            return path_parts[-2]
        return "Root"

    def count_files_recursive(self, api_url, params, target_extensions):
        """Counts total files to process for progress tracking."""
        count = 0
        try:
            response = requests.get(api_url, headers=self.get_headers(), params=params)
            print(f"[DEBUG] Count API Response Status: {response.status_code} for {api_url}")
            
            if response.status_code == 404:
                self.error = f"Repository or folder not found. Please check the URL."
                return 0
            elif response.status_code == 403:
                self.error = f"Access forbidden. Rate limit exceeded or private repo requires token."
                return 0
            elif response.status_code == 401:
                self.error = f"Authentication failed. Please check your GitHub token."
                return 0
                
            response.raise_for_status()
            items = response.json()
            
            # Check if response is an error message
            if isinstance(items, dict) and 'message' in items:
                self.error = f"GitHub API Error: {items['message']}"
                print(f"[DEBUG] GitHub API Error: {items['message']}")
                return 0
                
        except requests.exceptions.RequestException as e:
            self.error = f"Network error while counting files: {str(e)}"
            print(f"[DEBUG] Network Error: {e}")
            return 0
        except Exception as e:
            self.error = f"Error counting files: {str(e)}"
            print(f"[DEBUG] Exception: {e}")
            return 0

        if isinstance(items, dict):
            items = [items]
        
        for item in items:
            if item['type'] == 'file':
                if any(item['name'].endswith(ext) for ext in target_extensions):
                    count += 1
                    print(f"[DEBUG] Found matching file: {item['name']}")
            elif item['type'] == 'dir':
                count += self.count_files_recursive(item['url'], params, target_extensions)
        
        return count

    def fetch_files_recursive(self, api_url, params, target_extensions, current_folder="Root"):
        """Recursively fetches files from the GitHub API."""
        self.status = f"Scanning: {api_url.split('/')[-1]}"
        try:
            response = requests.get(api_url, headers=self.get_headers(), params=params)
            print(f"[DEBUG] Fetch API Response Status: {response.status_code} for {api_url}")
            
            if response.status_code == 404:
                self.error = f"Repository or folder not found: {api_url.split('/')[-1]}"
                return
            elif response.status_code == 403:
                self.error = f"Access forbidden. Rate limit exceeded or private repo requires token."
                return
            elif response.status_code == 401:
                self.error = f"Authentication failed. Please check your GitHub token."
                return
                
            response.raise_for_status()
            items = response.json()
            
            # Check if response is an error message
            if isinstance(items, dict) and 'message' in items:
                self.error = f"GitHub API Error: {items['message']}"
                print(f"[DEBUG] GitHub API Error: {items['message']}")
                return
                
        except requests.exceptions.RequestException as e:
            self.error = f"Network error while fetching files: {str(e)}"
            print(f"[DEBUG] Network Error: {e}")
            return
        except Exception as e:
            self.error = f"Failed to fetch folder contents: {e}"
            print(f"[DEBUG] Exception: {e}")
            return

        if isinstance(items, dict):
            items = [items]
        
        if isinstance(items, list):
            items.sort(key=lambda x: self.natural_key(x['name']))
            print(f"[DEBUG] Found {len(items)} items in folder")

        for item in items:
            if item['type'] == 'file':
                if any(item['name'].endswith(ext) for ext in target_extensions):
                    folder_name = self.get_folder_name_from_path(item['path'])
                    self.process_file(item, folder_name)
            
            elif item['type'] == 'dir':
                self.fetch_files_recursive(item['url'], params, target_extensions, item['name'])

    def process_file(self, file_item, folder_name):
        """Downloads raw content and formats it beautifully in Word."""
        name = file_item['name']
        download_url = file_item['download_url']
        path = file_item['path']

        self.current_file = name
        self.status = f"Processing: {name}"
        
        try:
            content_resp = requests.get(download_url, headers=self.get_headers())
            content_resp.raise_for_status()
            code_content = content_resp.text
            
            doc = self.get_or_create_document(folder_name)
            
            doc.add_heading(name, level=1)
            
            p_path = doc.add_paragraph()
            p_path.paragraph_format.space_after = Pt(6)
            run_path = p_path.add_run(f"Full Path: {path}")
            run_path.italic = True
            run_path.font.name = 'Segoe UI'
            run_path.font.size = Pt(8)
            run_path.font.color.rgb = RGBColor(128, 128, 128)

            code_paragraph = doc.add_paragraph(code_content)
            
            code_paragraph.style = 'No Spacing'
            for run in code_paragraph.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            self._set_paragraph_shading(code_paragraph, "F2F2F2")
            
            code_paragraph.paragraph_format.left_indent = Inches(0.1)
            code_paragraph.paragraph_format.space_before = Pt(6)
            code_paragraph.paragraph_format.space_after = Pt(12)

            doc.add_page_break()
            self.processed_files += 1
            
        except Exception as e:
            self.error = f"Could not download file {name}: {e}"

    def save_all_documents(self):
        """Saves all folder documents to separate files."""
        saved_files = []
        if not self.folder_documents:
            return saved_files
            
        for folder_name, doc in self.folder_documents.items():
            safe_filename = self.sanitize_filename(folder_name)
            output_path = os.path.join(self.output_directory, f"{safe_filename}.docx")
            doc.save(output_path)
            saved_files.append({
                'folder': folder_name,
                'filename': f"{safe_filename}.docx",
                'path': output_path
            })
        
        return saved_files

    def run(self, folder_url, target_extensions=['.cpp', '.h', '.hpp']):
        """Main execution function."""
        try:
            self.source_url = folder_url
            self.status = "parsing_url"
            print(f"[DEBUG] Starting conversion for URL: {folder_url}")
            print(f"[DEBUG] Target extensions: {target_extensions}")
            
            api_url, params = self.parse_github_url(folder_url)
            print(f"[DEBUG] Parsed API URL: {api_url}")
            print(f"[DEBUG] Params: {params}")
            
            self.status = "counting_files"
            self.total_files = self.count_files_recursive(api_url, params, target_extensions)
            print(f"[DEBUG] Total files found: {self.total_files}")
            
            # Check if any error occurred during counting
            if self.error:
                self.status = "error"
                return []
            
            # Check if no files were found
            if self.total_files == 0:
                self.error = f"No files found with extensions {', '.join(target_extensions)}. Please check the URL and ensure the repository contains files with those extensions."
                self.status = "error"
                return []
            
            self.status = "processing"
            self.fetch_files_recursive(api_url, params, target_extensions)
            
            # Check if any error occurred during fetching
            if self.error:
                self.status = "error"
                return []
            
            self.status = "saving"
            saved_files = self.save_all_documents()
            
            # Check if no documents were created
            if len(saved_files) == 0:
                self.error = "No documents were generated. Files may have failed to download."
                self.status = "error"
                return []
            
            self.status = "completed"
            print(f"[DEBUG] Conversion completed. Generated {len(saved_files)} document(s).")
            return saved_files
            
        except ValueError as e:
            self.error = f"Invalid URL format: {str(e)}"
            self.status = "error"
            print(f"[DEBUG] ValueError: {e}")
            return []
        except Exception as e:
            self.error = str(e)
            self.status = "error"
            print(f"[DEBUG] Exception: {e}")
            return []


def process_job(job_id, url, token, extensions, output_dir):
    """Background job processor."""
    jobs[job_id]['status'] = 'processing'
    print(f"[JOB {job_id}] Starting processing...")
    
    agent = GitHubFolderAgent(
        output_directory=output_dir,
        github_token=token,
        job_id=job_id
    )
    
    jobs[job_id]['agent'] = agent
    
    try:
        saved_files = agent.run(url, extensions)
        jobs[job_id]['files'] = saved_files
        
        # Check if agent encountered an error
        if agent.error:
            jobs[job_id]['status'] = 'error'
            jobs[job_id]['error'] = agent.error
            print(f"[JOB {job_id}] Error: {agent.error}")
        elif len(saved_files) == 0:
            jobs[job_id]['status'] = 'error'
            jobs[job_id]['error'] = 'No documents were generated. Please check the URL and file extensions.'
            print(f"[JOB {job_id}] No files generated")
        else:
            jobs[job_id]['status'] = 'completed'
            print(f"[JOB {job_id}] Completed with {len(saved_files)} file(s)")
    except Exception as e:
        jobs[job_id]['status'] = 'error'
        jobs[job_id]['error'] = str(e)
        print(f"[JOB {job_id}] Exception: {e}")


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint."""
    return jsonify({'status': 'healthy', 'message': 'GitHub to DOCX API is running'})


@app.route('/api/convert', methods=['POST'])
def start_conversion():
    """Start a new conversion job."""
    data = request.json
    
    url = data.get('url', '')
    token = data.get('token', None)
    extensions = data.get('extensions', ['.cpp', '.h', '.hpp', '.py', '.js'])
    
    if not url:
        return jsonify({'error': 'GitHub URL is required'}), 400
    
    # Validate URL format
    if 'github.com' not in url:
        return jsonify({'error': 'Please provide a valid GitHub URL'}), 400
    
    job_id = str(uuid.uuid4())
    output_dir = os.path.join('output', job_id)
    os.makedirs(output_dir, exist_ok=True)
    
    jobs[job_id] = {
        'id': job_id,
        'status': 'queued',
        'url': url,
        'files': [],
        'error': None,
        'output_dir': output_dir,
        'created_at': datetime.datetime.now().isoformat()
    }
    
    # Start processing in background thread
    thread = threading.Thread(
        target=process_job,
        args=(job_id, url, token, extensions, output_dir)
    )
    thread.start()
    
    return jsonify({'job_id': job_id, 'status': 'queued'})


@app.route('/api/status/<job_id>', methods=['GET'])
def get_status(job_id):
    """Get the status of a conversion job."""
    if job_id not in jobs:
        return jsonify({'error': 'Job not found'}), 404
    
    job = jobs[job_id]
    agent = job.get('agent')
    
    response = {
        'id': job_id,
        'status': job['status'],
        'files': job.get('files', []),
        'error': job.get('error')
    }
    
    if agent:
        response['progress'] = {
            'processed': agent.processed_files,
            'total': agent.total_files,
            'current_file': agent.current_file,
            'detail_status': agent.status
        }
    
    return jsonify(response)


@app.route('/api/download/<job_id>/<filename>', methods=['GET'])
def download_file(job_id, filename):
    """Download a generated DOCX file."""
    if job_id not in jobs:
        return jsonify({'error': 'Job not found'}), 404
    
    job = jobs[job_id]
    file_path = os.path.join(job['output_dir'], filename)
    
    if not os.path.exists(file_path):
        return jsonify({'error': 'File not found'}), 404
    
    return send_file(
        file_path,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/api/cleanup/<job_id>', methods=['DELETE'])
def cleanup_job(job_id):
    """Clean up a completed job and its files."""
    if job_id not in jobs:
        return jsonify({'error': 'Job not found'}), 404
    
    job = jobs[job_id]
    output_dir = job.get('output_dir')
    
    if output_dir and os.path.exists(output_dir):
        shutil.rmtree(output_dir)
    
    del jobs[job_id]
    
    return jsonify({'message': 'Job cleaned up successfully'})


if __name__ == '__main__':
    print("🚀 GitHub to DOCX API Server")
    print("=" * 40)
    print("Server running at http://localhost:5000")
    print("=" * 40)
    app.run(debug=True, port=5000)
