import os
import requests
import re
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from urllib.parse import urlparse

class GitHubFolderAgent:
    def __init__(self, output_directory="Output_Reports", github_token=None):
        self.output_directory = output_directory
        self.github_token = github_token
        self.folder_documents = {}  # Dictionary to store documents per folder
        self.source_url = ""
        
        # Create output directory if it doesn't exist
        if not os.path.exists(self.output_directory):
            os.makedirs(self.output_directory)

    def _setup_document_style(self, doc):
        """Sets up a professional visual style for the Word document."""
        # 1. Configure Normal Text
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Segoe UI'  # Clean, modern sans-serif
        font.size = Pt(10.5)
        
        # 2. Configure Heading 1 (For File Names)
        h1 = doc.styles['Heading 1']
        h1.font.name = 'Segoe UI'
        h1.font.size = Pt(14)
        h1.font.color.rgb = RGBColor(31, 78, 121)  # Professional Dark Blue
        h1.font.bold = True
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(3)

    def _set_paragraph_shading(self, paragraph, color_hex):
        """
        Directly manipulates OXML to add background color/shading to a paragraph.
        This creates the 'Code Block' gray box effect.
        """
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)  # Hex Color
        paragraph.paragraph_format.element.get_or_add_pPr().append(shd)

    def get_headers(self):
        """Returns headers for GitHub API requests."""
        headers = {'Accept': 'application/vnd.github.v3+json'}
        if self.github_token:
            headers['Authorization'] = f'token {self.github_token}'
        return headers

    def parse_github_url(self, url):
        """Converts a browser URL to a GitHub API URL."""
        parsed = urlparse(url)
        path_parts = parsed.path.strip('/').split('/')
        
        if len(path_parts) < 2:
            raise ValueError("Invalid GitHub URL")

        owner = path_parts[0]
        repo = path_parts[1]
        
        api_url = f"https://api.github.com/repos/{owner}/{repo}/contents"
        params = {}

        if len(path_parts) > 3 and path_parts[2] == 'tree':
            branch = path_parts[3]
            subpath = "/".join(path_parts[4:])
            api_url = f"{api_url}/{subpath}"
            params['ref'] = branch
            
        return api_url, params

    def natural_key(self, string_):
        """Returns a key for natural sorting."""
        return [int(s) if s.isdigit() else s.lower() for s in re.split(r'(\d+)', string_)]

    def sanitize_filename(self, name):
        """Sanitizes folder name to be used as a valid filename."""
        # Remove or replace invalid characters for Windows filenames
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            name = name.replace(char, '_')
        return name.strip()

    def get_or_create_document(self, folder_name):
        """Gets existing document for folder or creates a new one."""
        if folder_name not in self.folder_documents:
            doc = Document()
            self._setup_document_style(doc)
            
            # --- Create a Professional Title Page ---
            title = doc.add_heading(f'{folder_name}', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title.runs[0].font.name = 'Segoe UI'
            title.runs[0].font.color.rgb = RGBColor(16, 54, 92)  # Dark Blue

            # Add Date and URL info
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.add_run(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n").bold = True
            p.add_run(f"Source: {self.source_url}")
            doc.add_page_break()
            
            self.folder_documents[folder_name] = doc
            print(f"[Agent] Created new document for folder: {folder_name}")
        
        return self.folder_documents[folder_name]

    def get_folder_name_from_path(self, file_path):
        """Extracts the immediate parent folder name from a file path."""
        path_parts = file_path.split('/')
        if len(path_parts) > 1:
            return path_parts[-2]  # Return the parent folder name
        return "Root"  # If file is in root directory

    def fetch_files_recursive(self, api_url, params, target_extensions, current_folder="Root"):
        """Recursively fetches files from the GitHub API."""
        print(f"[Agent] Checking: {api_url}")
        try:
            response = requests.get(api_url, headers=self.get_headers(), params=params)
            response.raise_for_status()
            items = response.json()
        except Exception as e:
            print(f"[Error] Failed to fetch folder contents: {e}")
            return

        if isinstance(items, dict):
            items = [items]
        
        if isinstance(items, list):
            items.sort(key=lambda x: self.natural_key(x['name']))

        for item in items:
            if item['type'] == 'file':
                if any(item['name'].endswith(ext) for ext in target_extensions):
                    # Determine the folder this file belongs to
                    folder_name = self.get_folder_name_from_path(item['path'])
                    self.process_file(item, folder_name)
            
            elif item['type'] == 'dir':
                # Pass the directory name as the current folder
                self.fetch_files_recursive(item['url'], params, target_extensions, item['name'])

    def process_file(self, file_item, folder_name):
        """Downloads raw content and formats it beautifully in Word."""
        name = file_item['name']
        download_url = file_item['download_url']
        path = file_item['path']

        print(f"[Agent] Downloading: {name} -> Folder: {folder_name}")
        
        try:
            content_resp = requests.get(download_url, headers=self.get_headers())
            content_resp.raise_for_status()
            code_content = content_resp.text
            
            # Get or create the document for this folder
            doc = self.get_or_create_document(folder_name)
            
            # --- STYLING IMPROVEMENTS HERE ---
            
            # 1. Filename Heading
            doc.add_heading(name, level=1)
            
            # 2. File Path (Subtle, gray, italic)
            p_path = doc.add_paragraph()
            p_path.paragraph_format.space_after = Pt(6)
            run_path = p_path.add_run(f"Full Path: {path}")
            run_path.italic = True
            run_path.font.name = 'Segoe UI'
            run_path.font.size = Pt(8)
            run_path.font.color.rgb = RGBColor(128, 128, 128)  # Grey

            # 3. The Code Block (The "Pretty" Part)
            # We create a new paragraph for the code
            code_paragraph = doc.add_paragraph(code_content)
            
            # 3a. Format the text (Monospace font, smaller size)
            code_paragraph.style = 'No Spacing'
            for run in code_paragraph.runs:
                run.font.name = 'Consolas'  # Best font for code
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # 3b. Add Background Shading (Light Gray #F2F2F2)
            self._set_paragraph_shading(code_paragraph, "F2F2F2")
            
            # 3c. Adjust spacing around the code block
            code_paragraph.paragraph_format.left_indent = Inches(0.1)
            code_paragraph.paragraph_format.space_before = Pt(6)
            code_paragraph.paragraph_format.space_after = Pt(12)

            doc.add_page_break()
            
        except Exception as e:
            print(f"[Error] Could not download file {name}: {e}")

    def save_all_documents(self):
        """Saves all folder documents to separate files."""
        if not self.folder_documents:
            print("[Warning] No documents to save!")
            return
            
        for folder_name, doc in self.folder_documents.items():
            safe_filename = self.sanitize_filename(folder_name)
            output_path = os.path.join(self.output_directory, f"{safe_filename}.docx")
            doc.save(output_path)
            print(f"[Success] Saved: {output_path}")
        
        print(f"\n[Complete] Total {len(self.folder_documents)} document(s) saved in '{self.output_directory}' folder.")

    def run(self, folder_url, target_extensions=['.cpp', '.h', '.hpp']):
        """Main execution function."""
        try:
            self.source_url = folder_url
            
            # --- Start Processing ---
            api_url, params = self.parse_github_url(folder_url)
            print(f"[Agent] Targeting API: {api_url}")
            
            self.fetch_files_recursive(api_url, params, target_extensions)
            
            # Save all documents
            self.save_all_documents()
            
        except Exception as e:
            print(f"[Fatal Error] {e}")


# --- CONFIGURATION SECTION ---
if __name__ == "__main__":
    # 1. PASTE YOUR GITHUB SUBFOLDER LINK HERE
    TARGET_URL = "https://github.com/YourName/YourRepo/tree/main/YourSubfolder"

    # 2. (Optional) If repo is PRIVATE, put your Personal Access Token here:
    TOKEN = None

    # 3. Output directory where all folder-wise .docx files will be saved
    OUTPUT_DIR = "Output_Reports"

    agent = GitHubFolderAgent(output_directory=OUTPUT_DIR, github_token=TOKEN)
    
    # 4. Run
    agent.run(TARGET_URL, target_extensions=['.cpp', '.h', '.hpp', '.py', '.js'])
